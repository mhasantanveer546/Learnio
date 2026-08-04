from docx import Document


def extract_text_from_docx(filepath):
    doc = Document(filepath)

    # Paragraphs AND tables — a lecture note doc often has both, and
    # skipping tables silently would lose real content (e.g. formula
    # sheets, comparison tables).
    parts = [p.text for p in doc.paragraphs if p.text.strip()]

    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            if row_text.strip(" |"):
                parts.append(row_text)

    extracted = "\n".join(parts).strip()

    if not extracted:
        raise ValueError("No extractable text found in this document.")

    return extracted